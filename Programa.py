import sys
from xml.etree.ElementTree import parse
from numpy import empty, string_
from numpy.core.fromnumeric import choose
from numpy.lib.twodim_base import diag
import pandas as pd
import re
import docx
from PyQt5 import uic, QtWidgets, QtGui
from PyQt5.QtCore import QDate
from pandas.core.frame import DataFrame


def gerarWordRelTrimestral():

    caminhoArquivo = pickArchive()
    print(caminhoArquivo[0])

    dataframe = pd.read_excel(
        r""+caminhoArquivo[0], sheet_name="Teste")

    listaUnidades = []
    listaArea = []

    doc = docx.Document()

    # Carrega informações de Unidades
    for index, row in dataframe.iterrows():
        listaUnidades.append(row["Unidade"])

    # Exclui repetidos
    listaUnidades = list(dict.fromkeys(listaUnidades))

    for unidade in listaUnidades:

        df_maskUnidade = dataframe["Unidade"] == unidade
        filtradoUnidade = dataframe[df_maskUnidade]
        listaServicos = []
        # Carrega informações de Serviço
        for index, row in filtradoUnidade.iterrows():
            listaServicos.append(row["Serviço"])

        # Exclui repetidos
        listaServicos = list(dict.fromkeys(listaServicos))

        doc.add_paragraph().add_run(str(unidade))

        for servicos in listaServicos:
            texto = servicos + " ("
            listaArea = []
            df_maskServico = filtradoUnidade["Serviço"] == servicos
            filtradoServico = filtradoUnidade[df_maskServico]
            for index, row in filtradoServico.iterrows():
                listaArea.append(row["Área"])

        # FUNCIONANDO

            if len(listaArea) == 1:
                texto = texto + str(listaArea[0])
            else:
                for i in range(len(listaArea)):
                    texto = texto + str(listaArea[i])
                    if (i+1) != len(listaArea):
                        texto = texto + str("/")
            texto = texto + str(")")
            doc.add_paragraph().add_run(re.sub(' +', ' ', texto))

    caminho = str(chooseSavePath())
    if not caminho:
        doc.save(r""+r"C:\Users\pedror\Documents\Text.docx")
    else:
        doc.save(r""+caminho+"/Text.docx")

    msg = QtWidgets.QMessageBox()
    msg.information(
        None, "Sucesso", "Realizado com sucesso", QtWidgets.QMessageBox.Ok)


def analiseRel20():

    caminhoArquivo = pickArchive()

    dataFiltro = formulario.dataFiltroFS.text()

    dataframe = pd.read_excel(
        r""+caminhoArquivo[0], sheet_name="Novo Relatório")

    filtradofinal = dataframe.query(
        'DT_EXECUCAO == "" and STATUS == "OK" and DT_PROGRAMACAO <= "' + dataFiltro + '"')
    
    filtradofinal = filtradofinal.drop(columns=['PASTA_TH', 'PONDERACAO_PREVISTO', 'SEMANA_EXEC', 'SEMANA_PROG'])

    # Remover colunas PASTA_TH, PONDERACAO_PREVISTO, SEMANA_EXEC, SEMANA_PROG

    listaDisciplinas = filtradofinal['DISC_NOME'].unique().tolist()

    path = chooseSavePath()

    for i in listaDisciplinas:
        filtradofinal.query('DISC_NOME == "' + i + '"').to_excel(r"" + path +
                                                                 "/"+i+" até "+dataFiltro.replace("/", "-")+".xlsx", index=False)

    msg = QtWidgets.QMessageBox()
    msg.information(None, "Sucesso", "Realizado com sucesso",
                    QtWidgets.QMessageBox.Ok)


def importParaSisepc():

    dataDate = formulario.dataDate.text()

    colunasCriacao = ['numero', 'revisao', 'situacao', 'criterioMedicao', 'descricao', 'subcontrato',
                      'unidadeProcesso', 'areaAplicacao', 'produtoFabricao', 'obs', 'qdPrev', 'qtReal', 'unidadeMedida']
    colunasVinculo = ['numeroFS', 'criterioMedicao', 'sigla',
                      'IDatividade', 'IDatividade2', 'Responsavel', 'UA']
    colunasAvanco = ['IDatividade', 'numeroFS', 'WBS', 'Nivel', 'Sigla', 'descricao',
                     'col7', 'col8', 'col9', 'col10', 'col11', 'dataExec', 'col13', 'Avanco', 'col15', 'col16', 'col17']

    dfCriarFS = pd.DataFrame(columns=colunasCriacao)
    auxCriarFS = pd.DataFrame(columns=colunasCriacao)
    dfVinculo = pd.DataFrame(columns=colunasVinculo)
    dfAvanco = pd.DataFrame(columns=colunasAvanco)
    auxAvanco = pd.DataFrame(columns=colunasAvanco)

    # BLOCO DE LEITURA DO ARQUIVO .CSV

    caminho = pickArchive()
    df = pd.read_csv(r""+caminho[0], header=1, skipfooter=1, quotechar='"',
                     sep=";", encoding="iso8859-1", engine="python")

    df = df.fillna('')
    df = df.rename(columns={'Activity ID': 'ID', 'Activity Status': 'Status', 'Activity Name': 'Name',
                   'BL1 Activity % Complete': 'BL1Percent', 'Activity % Complete': 'NewPercent', 'Actual Start': 'Start', 'Actual Finish': 'Finish'})

    df['NewPercent'] = df['NewPercent'].str.rstrip('%')
    df['BL1Percent'] = df['BL1Percent'].str.rstrip('%')
    df['ID'] = df['ID'].str.lstrip()
    df['NewPercent'] = pd.to_numeric(df['NewPercent'], downcast="float")
    df['BL1Percent'] = pd.to_numeric(df['BL1Percent'], downcast="float")

    # AVANÇO INICIAL ----------- OK

    av_parcialNovo = df.query(
        '(BL1Percent < NewPercent) and (Status == "In Progress") and (BL1Percent == 0)')

    # BLOCO - CRIAÇÃO DE FS DO PRIMEIRO FILTRO

    auxCriarFS['numero'] = av_parcialNovo['ID']
    auxCriarFS['revisao'] = auxCriarFS['revisao'].fillna('0')
    auxCriarFS['situacao'] = auxCriarFS['situacao'].fillna('OK')
    auxCriarFS['criterioMedicao'] = auxCriarFS['criterioMedicao'].fillna(
        'AVANÇO CRONOGRAMA')
    auxCriarFS['descricao'] = auxCriarFS['descricao'].fillna('')
    auxCriarFS['subcontrato'] = auxCriarFS['subcontrato'].fillna('')
    auxCriarFS['unidadeProcesso'] = auxCriarFS['unidadeProcesso'].fillna(
        'GERAL')
    auxCriarFS['areaAplicacao'] = auxCriarFS['areaAplicacao'].fillna('')
    auxCriarFS['produtoFabricao'] = auxCriarFS['produtoFabricao'].fillna('')
    auxCriarFS['obs'] = auxCriarFS['obs'].fillna('')
    auxCriarFS['qdPrev'] = auxCriarFS['qdPrev'].fillna('1')
    auxCriarFS['qtReal'] = auxCriarFS['qtReal'].fillna('0')
    auxCriarFS['unidadeMedida'] = auxCriarFS['unidadeMedida'].fillna('un')

    dfCriarFS = dfCriarFS.append(
        auxCriarFS, ignore_index=False, verify_integrity=False, sort=False)

    # AVANÇO INICIAL E FINAL ----------- OK

    av_totalNovo = df.query(
        '(BL1Percent < NewPercent) and (Status == "Completed") and (BL1Percent == 0)')

    # BLOCO - CRIAÇÃO DE FS DO SEGUNDO FILTRO

    auxCriarFS = pd.DataFrame(columns=colunasCriacao)
    auxCriarFS['numero'] = av_totalNovo['ID']
    auxCriarFS['revisao'] = auxCriarFS['revisao'].fillna('0')
    auxCriarFS['situacao'] = auxCriarFS['situacao'].fillna('OK')
    auxCriarFS['criterioMedicao'] = auxCriarFS['criterioMedicao'].fillna(
        'AVANÇO CRONOGRAMA')
    auxCriarFS['descricao'] = auxCriarFS['descricao'].fillna('')
    auxCriarFS['subcontrato'] = auxCriarFS['subcontrato'].fillna('')
    auxCriarFS['unidadeProcesso'] = auxCriarFS['unidadeProcesso'].fillna(
        'GERAL')
    auxCriarFS['areaAplicacao'] = auxCriarFS['areaAplicacao'].fillna('')
    auxCriarFS['produtoFabricao'] = auxCriarFS['produtoFabricao'].fillna('')
    auxCriarFS['obs'] = auxCriarFS['obs'].fillna('')
    auxCriarFS['qdPrev'] = auxCriarFS['qdPrev'].fillna('1')
    auxCriarFS['qtReal'] = auxCriarFS['qtReal'].fillna('0')
    auxCriarFS['unidadeMedida'] = auxCriarFS['unidadeMedida'].fillna('un')

    dfCriarFS = dfCriarFS.append(
        auxCriarFS, ignore_index=False, verify_integrity=False, sort=False)

    # BLOCO CRIACAO DO DATAFRAME DE VINCULO

    dfVinculo['numeroFS'] = dfCriarFS['numero']
    dfVinculo['sigla'] = dfVinculo['sigla'].fillna('1.')
    dfVinculo = dfVinculo.fillna("")
    dfVinculo['IDatividade'] = dfVinculo['numeroFS']

    # A PARTIR DAQUI NÃO PRECISA MAIS CRIAR FS, SÓ AVANÇO

    #             AVANÇO PRIMEIRO FILTRO
    if av_parcialNovo.empty:
        print("pular filtro 1")
    else:
        auxAvanco['IDatividade'] = av_parcialNovo['ID']
        auxAvanco['numeroFS'] = av_parcialNovo['ID']
        auxAvanco['WBS'] = auxAvanco['WBS'].fillna('1.1.')
        auxAvanco['Nivel'] = auxAvanco['Nivel'].fillna('2')
        auxAvanco['Sigla'] = auxAvanco['Sigla'].fillna('1.1.')
        auxAvanco['descricao'] = auxAvanco['descricao'].fillna(
            'AVANÇO CRONOGRAMA')
        auxAvanco['dataExec'] = av_parcialNovo['Start']
        auxAvanco['Avanco'] = av_parcialNovo['NewPercent']
        auxAvanco = auxAvanco.fillna('')

        dfAvanco = dfAvanco.append(
            auxAvanco, ignore_index=False, verify_integrity=False, sort=False)

    #             AVANÇO SEGUNDO FILTRO

    auxAvanco = pd.DataFrame(columns=colunasAvanco)

    if av_totalNovo.empty:
        print("pular filtro 2")
    else:
        auxAvanco['IDatividade'] = av_totalNovo['ID']
        auxAvanco['numeroFS'] = av_totalNovo['ID']
        auxAvanco['WBS'] = auxAvanco['WBS'].fillna('1.1.')
        auxAvanco['Nivel'] = auxAvanco['Nivel'].fillna('2')
        auxAvanco['Sigla'] = auxAvanco['Sigla'].fillna('1.1.')
        auxAvanco['descricao'] = auxAvanco['descricao'].fillna(
            'AVANÇO CRONOGRAMA')
        auxAvanco['dataExec'] = av_totalNovo['Start']
        auxAvanco['Avanco'] = auxAvanco['Avanco'].fillna('1')
        auxAvanco = auxAvanco.fillna('')

        dfAvanco = dfAvanco.append(
            auxAvanco, ignore_index=False, verify_integrity=False, sort=False)
        auxAvanco = pd.DataFrame(columns=colunasAvanco)

        auxAvanco['IDatividade'] = av_totalNovo['ID']
        auxAvanco['numeroFS'] = av_totalNovo['ID']
        auxAvanco['WBS'] = auxAvanco['WBS'].fillna('1.1.')
        auxAvanco['Nivel'] = auxAvanco['Nivel'].fillna('2')
        auxAvanco['Sigla'] = auxAvanco['Sigla'].fillna('1.1.')
        auxAvanco['descricao'] = auxAvanco['descricao'].fillna(
            'AVANÇO CRONOGRAMA')
        auxAvanco['dataExec'] = av_totalNovo['Finish']
        auxAvanco['Avanco'] = auxAvanco['Avanco'].fillna('100')
        auxAvanco = auxAvanco.fillna('')

        dfAvanco = dfAvanco.append(
            auxAvanco, ignore_index=False, verify_integrity=False, sort=False)

    # AVANÇO PARCIAL (NO DATA DATE) - TERCEIRO FILTRO
    av_parcial = df.query(
        '(BL1Percent < NewPercent) and (Status == "In Progress") and (BL1Percent != 0)')

    auxAvanco = pd.DataFrame(columns=colunasAvanco)

    if av_parcial.empty:
        print("pular filtro 3")
    else:
        auxAvanco['IDatividade'] = av_parcial['ID']
        auxAvanco['numeroFS'] = av_parcial['ID']
        auxAvanco['WBS'] = auxAvanco['WBS'].fillna('1.1.')
        auxAvanco['Nivel'] = auxAvanco['Nivel'].fillna('2')
        auxAvanco['Sigla'] = auxAvanco['Sigla'].fillna('1.1.')
        auxAvanco['descricao'] = auxAvanco['descricao'].fillna(
            'AVANÇO CRONOGRAMA')
        auxAvanco['dataExec'] = dataDate
        auxAvanco['Avanco'] = av_parcial['NewPercent']
        auxAvanco = auxAvanco.fillna('')

        dfAvanco = dfAvanco.append(
            auxAvanco, ignore_index=False, verify_integrity=False, sort=False)

    # AVANÇO FINAL (NO FINISH)

    av_total = df.query(
        '(BL1Percent < NewPercent) and (Status == "Completed") and (BL1Percent > 0)')

    auxAvanco = pd.DataFrame(columns=colunasAvanco)

    if av_total.empty:
        print("pular filtro 4")
    else:
        auxAvanco['IDatividade'] = av_total['ID']
        auxAvanco['numeroFS'] = av_total['ID']
        auxAvanco['WBS'] = auxAvanco['WBS'].fillna('1.1.')
        auxAvanco['Nivel'] = auxAvanco['Nivel'].fillna('2')
        auxAvanco['Sigla'] = auxAvanco['Sigla'].fillna('1.1.')
        auxAvanco['descricao'] = auxAvanco['descricao'].fillna(
            'AVANÇO CRONOGRAMA')
        auxAvanco['dataExec'] = av_total['Finish']
        auxAvanco['Avanco'] = auxAvanco['Avanco'].fillna('100')
        auxAvanco = auxAvanco.fillna('')

        dfAvanco = dfAvanco.append(
            auxAvanco, ignore_index=False, verify_integrity=False, sort=False)

    dfAvanco = dfAvanco.query('dataExec != ""')

    path = chooseSavePath()

    salvar = pd.ExcelWriter(
        r""+path+'\Importacao.xlsx')
    dfCriarFS.to_excel(salvar, index=False, sheet_name="Cria FS")
    dfVinculo.to_excel(salvar, index=False, sheet_name="Vinculo FS")
    dfAvanco.to_excel(salvar, index=False, sheet_name="Avanços")

    salvar.save()


# escolha .xls* ou .csv do arquivo pra importar
def pickArchive():
    dialog = QtWidgets.QFileDialog()
    folder_path = dialog.getOpenFileName(
        None, "Select Excel file to import", "", "Excel (*.csv *.xls *.xlsx)")
    return folder_path

# do arquivo .docx


def chooseSavePath():
    dialog = QtWidgets.QFileDialog()
    folder_path = str(dialog.getExistingDirectory(
        None, "Escolha a porra da pasta onde salvar..."))
    return folder_path


app = QtWidgets.QApplication([])
formulario = uic.loadUi("telaPrincipal.ui")

# Rotas dos botões
formulario.btnRodar.clicked.connect(gerarWordRelTrimestral)
formulario.btnAnalisar.clicked.connect(analiseRel20)
formulario.btnImportSisepc.clicked.connect(importParaSisepc)

# Alterando o campo de data para o Current date
now = QDate.currentDate()
formulario.dataFiltroFS.setDate(now)

formulario.show()
app.exec()
